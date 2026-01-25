from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def extract_text_from_docx(file_path):
    """
    Extracts full text from a DOCX file.
    """
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)


from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def keep_heading_with_content(paragraph):
    """Rule A: Keep with Next (no orphan headings)"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    keepNext = pPr.find(qn('w:keepNext'))
    if keepNext is None:
        keepNext = OxmlElement('w:keepNext')
        pPr.append(keepNext)

def prevent_widow(paragraph):
    """Rule C: Prevent single-line section endings"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    widow = pPr.find(qn('w:widowControl'))
    if widow is None:
        widow = OxmlElement('w:widowControl')
        pPr.append(widow)

def apply_journal_formatting(input_path, output_path, analysis_data):
    """
    Rebuilds the DOCX with IEEE-like formatting, enforcing DocPhysics rules.
    """
    doc = Document(input_path)
    
    # 1. Prepare Header Mapping for Standardization
    header_map = {}
    if analysis_data.get("sections"):
        for section in analysis_data["sections"]:
            original = section.get("original_header", "").strip().lower()
            standardized = section.get("standardized_header", "").strip()
            if original and standardized:
                header_map[original] = standardized

    # 2. Base Styles (Times New Roman)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Enable global settings
    doc.settings.odd_and_even_pages_header_footer = False 

    # 3. Iterate paragraphs to apply "DocPhysics" rules
    keywords_inserted = False
    
    for i, para in enumerate(doc.paragraphs):
        # --- Rule: Standardization ---
        # Check if this paragraph is a header that needs replacing
        # We check both style name and if it matches our map
        text_clean = para.text.strip().lower()
        if text_clean in header_map:
            # It's a match! Replace text.
            para.text = header_map[text_clean]
            # Ensure it is styled as a heading if not already (heuristic)
            if not para.style.name.startswith("Heading"):
                para.style = doc.styles['Heading 1']

        # --- Rule: Typography & Alignment ---
        # Apply Justified alignment to normal paragraphs that are not headers or empty
        if not para.style.name.startswith("Heading") and para.text.strip():
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Apply standard pagination rules to ALL paragraphs
        # widow_control is a property on paragraph_format, but we'll use the OXML helper for robustness if needed, 
        # or just set the property. python-docx supports it directly.
        para.paragraph_format.widow_control = True

        # --- Rule: Metadata Injection ---
        # A. Insert Keywords after the first paragraph (index 0)
        # We assume para[0] is the title or banner. We do NOT touch it.
        if i == 0:
            # Ensure Title is centered and bold for the first paragraph
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(16) # Make title larger

            if analysis_data.get("keywords") and not keywords_inserted:
                kw_text = "Keywords: " + ", ".join(analysis_data["keywords"])
                # Add keywords after the first paragraph
                if len(doc.paragraphs) > 1:
                    # We can't easily insert *after* current in iteration without messing up, 
                    # but doc.paragraphs is a list snapshot? No, it's dynamic.
                    # Safest is to insert 'before' the *next* one if we were iterating differently,
                    # but here we can just append to doc if we were building new, but we are editing in place.
                    # Ideally we want it after para 0.
                    # Let's insert it before para 1?
                    pass # logic below handles insertion
                
                # We need to insert it *after* the current paragraph (which is i=0).
                # The python-docx API `insert_paragraph_before` is on the *next* paragraph.
                # If there is a next paragraph:
                try:
                    p = doc.paragraphs[i+1].insert_paragraph_before(kw_text)
                    p.style = doc.styles['Normal']
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_after = Pt(12) # Add some space
                    # p.runs[0] might not exist yet if we just set text on create?
                    # insert_paragraph_before(text) creates a run.
                    if p.runs:
                        p.runs[0].font.italic = True
                        p.runs[0].font.name = 'Times New Roman'
                        p.runs[0].font.size = Pt(12)
                    keywords_inserted = True
                except IndexError:
                    # Only 1 paragraph implementation
                    p = doc.add_paragraph(kw_text)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    keywords_inserted = True
            continue

        # --- Rule: Heading Physics ---
        # B. Pagination Fixes for Headers
        if para.style.name.startswith("Heading"):
             # Reinforce Keep with Next
             keep_heading_with_content(para)
             para.paragraph_format.keep_with_next = True
             para.paragraph_format.keep_together = True
             # Ensure headers are not justified, usually Left or Center
             # Let's stick to Left for standard IEEE-like, or keep original.
             # Forced Justify looks bad on short headers.
             if para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                 para.alignment = WD_ALIGN_PARAGRAPH.LEFT
             
             # Spacing optimization for headers
             para.paragraph_format.space_before = Pt(12)
             para.paragraph_format.space_after = Pt(6)

    # 4. Add Footer Verification
    section = doc.sections[0]
    
    # Set Margins (Standard 1 inch)
    # 914400 EMUs = 1 inch
    section.top_margin = 914400
    section.bottom_margin = 914400
    section.left_margin = 914400
    section.right_margin = 914400

    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.text = "Formatted by ResearchMate AI • Final Word"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer_para.runs:
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.name = 'Times New Roman' # Ensure footer font consistency
        
    doc.save(output_path)
    return output_path
